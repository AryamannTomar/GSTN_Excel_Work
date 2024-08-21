import tkinter as tk
from tkinter import filedialog
import pandas as pd
import pickle
from docx import Document
from docx.shared import Pt
import numpy as np
import os
import docx2pdf
import threading
from tqdm import tqdm

class ExcelUploaderApp:
    def __init__(self, root):
        self.root = root
        self.root.geometry("600x300")
        self.root.title("Excel File Uploader")
        
        self.upload_button1 = tk.Button(root, text="Upload File 1", command=self.upload_file1)
        self.upload_button1.pack(pady=10)
        
        self.result_label1 = tk.Label(root, text="")
        self.result_label1.pack(pady=5)
        
        self.upload_button2 = tk.Button(root, text="Upload File 2", command=self.upload_file2)
        self.upload_button2.pack(pady=10)
        
        self.result_label2 = tk.Label(root, text="")
        self.result_label2.pack(pady=5)

        self.submit_button = tk.Button(root, text="Submit Files", command=self.submit_files)
        self.submit_button.pack(pady=20)

        self.message_label = tk.Label(root, text="")
        self.message_label.pack(pady=5)

        self.file_path1 = None
        self.file_path2 = None
        self.lenx=None
        self.leny=None

    def process_document(self, fname, vals, f, m):
        doc = Document(fname)
        all_tables = []
        for table in doc.tables:
            rows = len(table.rows)
            cols = len(table.columns)
            data = np.empty((rows, cols), dtype=object)
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    data[i, j] = cell.text.strip()
            all_tables.append(data)

        for i, j in vals:
            self.update_document(i[0], i[1], j, doc.tables[i[2]])

        output_docx = f"{f}_Updated_Vals_{m}.docx"
        doc.save(output_docx)
        # docx2pdf.convert(output_docx, f"{f}_Updated_{m}.pdf")
        os.remove(output_docx)

    def update_document(self, row, col, new_value, table):
        if row >= len(table.rows) or col >= len(table.columns):
            print("Warning: The specified cell does not exist in the table.")
            return    
        cell = table.rows[row].cells[col]
        cell.text = str(new_value)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name, run.font.size = 'Roboto', Pt(7.5)

    def upload_file1(self):
        self.file_path1 = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls;*.csv")])
        if self.file_path1:
            try:
                df = pd.read_csv(self.file_path1)
                self.lenx=df.shape[0]
                self.result_label1.config(text=f" {self.file_path1.split('/')[-1]} file uploaded successfully. Shape: {df.shape}")
            except Exception as e:
                self.result_label1.config(text=f"Error: {str(e)}")
                self.file_path1 = None

    def upload_file2(self):
        self.file_path2 = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls;*.csv")])
        if self.file_path2:
            try:
                df = pd.read_csv(self.file_path2)
                self.leny=df.shape[0]
                self.result_label2.config(text=f" {self.file_path2.split('/')[-1]} file uploaded successfully. Shape: {df.shape}")
            except Exception as e:
                self.result_label2.config(text=f"Error: {str(e)}")
                self.file_path2 = None

    def process_document(self, fname, vals, f, m):
        doc = Document(fname)
        all_tables = []
        for table in doc.tables:
            rows = len(table.rows)
            cols = len(table.columns)
            data = np.empty((rows, cols), dtype=object)
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    data[i, j] = cell.text.strip()
            all_tables.append(data)

        for i, j in vals:
            self.update_document(i[0], i[1], j, doc.tables[i[2]])

        output_docx = f"{f}_Updated_Vals_{m}.docx"
        doc.save(output_docx)
        # docx2pdf.convert(output_docx, f"{f}_Updated_{m}.pdf")
        os.remove(output_docx)

    def update_document(self, row, col, new_value, table):
        if row >= len(table.rows) or col >= len(table.columns):
            print("Warning: The specified cell does not exist in the table.")
            return    
        cell = table.rows[row].cells[col]
        cell.text = str(new_value)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.name, run.font.size = 'Roboto', Pt(7.5)

    def submit_files(self):
        if self.file_path1 and self.file_path2 and self.lenx==self.leny:
            self.submit_button.config(state=tk.DISABLED)  
            
            def submit_process():
                for m in tqdm(range(self.lenx), desc='Processing', unit='iteration', file=None, leave=True):
                    l = pd.concat([pd.read_csv(self.file_path1), pd.read_csv(self.file_path2)], axis=0, ignore_index=True).to_dict('records')

                    with open('final_gstr3b_params.pkl', 'rb') as f:
                        x1 = pickle.load(f)
                    gstr3b = {i[1]:i[3] for i in x1}

                    with open('final_gstr1_params.pkl', 'rb') as f:
                        x2 = pickle.load(f)
                    gstr1 = {i[1]:i[3] for i in x2}

                    for m_index, x in enumerate(l):
                        if m_index == m:
                            d_ = {}
                            d1_ = {}
                            for gstr, d in [(gstr3b, d_), (gstr1, d1_)]:
                                for i in x.keys():
                                    l_ = [k for k, val in gstr.items() if val == i]
                                    if len(l_) != 0:
                                        d.update(dict(zip(l_, tuple([x[i]]) * len(l_))))
                                for i in set(gstr.keys()).difference(set(d.keys())):
                                    d[i] = '?'
                            self.process_document('gstr3b.docx', [(k, v) for k, v in d_.items()], 'gstr3b', m)
                            self.message_label.config(text=f'{'*'*((m+1)%10)} Completed {m+1}/{self.lenx} Rows {'*'*((m+1)%10)}')

                            self.process_document('gstr1.docx', [(k, v) for k, v in d1_.items()], 'gstr1', m)

                self.submit_button.config(state=tk.NORMAL) 

            submit_thread = threading.Thread(target=submit_process)
            submit_thread.start()

root = tk.Tk()
app = ExcelUploaderApp(root)
root.mainloop()